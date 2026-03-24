"""
Export Service - Document Export to Multiple Formats
Exports analysis results to Word (.docx), Markdown (.md), and HTML
Enhanced to properly format structured content into professional documents
"""

from typing import Dict, Any, List
import json
import logging
import re
from io import BytesIO

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting documents to various formats"""
    
    def __init__(self):
        # Check for python-docx availability
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.table import WD_TABLE_ALIGNMENT
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            logger.warning("python-docx not available. Word export will be limited.")
    
    # ==================== WORD EXPORT ====================
    
    def to_word(
        self,
        content: Any,
        analysis_type: str,
        project_id: str
    ) -> bytes:
        """
        Export document to Word format (.docx)
        """
        if not self.docx_available:
            return self._export_as_text(content, analysis_type).encode('utf-8')
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            doc = Document()
            
            # Set document margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Add cover page
            self._add_cover_page(doc, analysis_type, project_id)
            
            # Add table of contents placeholder
            self._add_toc_placeholder(doc)
            
            # Parse content
            content_dict = self._parse_content(content)
            
            # Add content based on type
            if analysis_type == 'brd':
                self._add_brd_content(doc, content_dict)
            elif analysis_type == 'frd':
                self._add_frd_content(doc, content_dict)
            elif analysis_type == 'user_stories':
                self._add_user_stories_content(doc, content_dict)
            elif analysis_type == 'test_cases':
                self._add_test_cases_content(doc, content_dict)
            elif analysis_type == 'migration_plan':
                self._add_migration_plan_content(doc, content_dict)
            elif analysis_type == 'reverse_eng':
                self._add_reverse_eng_content(doc, content_dict)
            elif analysis_type == 'tdd':
                self._add_tdd_content(doc, content_dict)
            elif analysis_type == 'db_analysis':
                self._add_db_analysis_content(doc, content_dict)
            else:
                self._add_generic_content(doc, content_dict)
            
            # Save to bytes
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            return file_stream.read()
            
        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            import traceback
            traceback.print_exc()
            return self._export_as_text(content, analysis_type).encode('utf-8')
    
    def _parse_content(self, content: Any) -> Dict[str, Any]:
        """Parse content from various formats"""
        if isinstance(content, str):
            try:
                content_dict = json.loads(content)
            except:
                content_dict = {'content': content}
        else:
            content_dict = content
        
        # Extract actual content if nested
        if 'content' in content_dict and isinstance(content_dict['content'], dict):
            content_dict = content_dict['content']
        
        return content_dict
    
    def _add_cover_page(self, doc, analysis_type: str, project_id: str):
        """Add a professional cover page"""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self._get_document_title(analysis_type))
        title_run.font.size = Pt(28)
        title_run.bold = True
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("Comprehensive Documentation")
        sub_run.font.size = Pt(16)
        sub_run.italic = True
        
        # More spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Project ID: {project_id}\n").font.size = Pt(11)
        meta.add_run(f"Document Type: {analysis_type.upper()}\n").font.size = Pt(11)
        meta.add_run(f"Generated by: Idea2Spec\n").font.size = Pt(11)
        
        # Page break
        doc.add_page_break()
    
    def _add_toc_placeholder(self, doc):
        """Add table of contents placeholder"""
        doc.add_heading('Table of Contents', 1)
        toc = doc.add_paragraph()
        toc.add_run('[Table of Contents - Update after opening in Word]')
        toc.add_run('\nRight-click and select "Update Field" to generate TOC')
        doc.add_page_break()
    
    def _format_text_content(self, text: str) -> str:
        """Format text content, handling special characters"""
        if not isinstance(text, str):
            text = str(text)
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()
    
    def _add_paragraph_text(self, doc, text: str, style: str = None):
        """Add a paragraph with properly formatted text"""
        text = self._format_text_content(text)
        if '\n\n' in text:
            # Multiple paragraphs
            for para_text in text.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip(), style=style)
        else:
            if text.strip():
                p = doc.add_paragraph(text.strip(), style=style)
    
    def _add_key_value_paragraph(self, doc, key: str, value: Any, bold_key: bool = True):
        """Add a key-value pair as a paragraph"""
        from docx.shared import Pt
        
        p = doc.add_paragraph()
        key_run = p.add_run(f"{key}: ")
        if bold_key:
            key_run.bold = True
        
        if isinstance(value, list):
            value_text = ", ".join(str(v) for v in value)
        elif isinstance(value, dict):
            value_text = ", ".join(f"{k}: {v}" for k, v in value.items())
        else:
            value_text = str(value)
        
        p.add_run(value_text)
    
    def _add_bullet_list(self, doc, items: List, level: int = 0):
        """Add a bullet list"""
        for item in items:
            if isinstance(item, dict):
                # Format dict items
                for k, v in item.items():
                    if k not in ['raw_output', 'error', 'timestamp']:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{self._format_key(k)}: ").bold = True
                        p.add_run(str(v) if not isinstance(v, (dict, list)) else self._format_complex_value(v))
            elif isinstance(item, str):
                doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(str(item), style='List Bullet')
    
    def _add_numbered_list(self, doc, items: List):
        """Add a numbered list"""
        for item in items:
            if isinstance(item, dict):
                text = self._format_dict_as_text(item)
                doc.add_paragraph(text, style='List Number')
            else:
                doc.add_paragraph(str(item), style='List Number')
    
    def _format_key(self, key: str) -> str:
        """Format a key for display"""
        return key.replace('_', ' ').title()
    
    def _format_dict_as_text(self, d: Dict) -> str:
        """Format a dictionary as readable text"""
        parts = []
        for k, v in d.items():
            if k not in ['raw_output', 'error', 'timestamp']:
                if isinstance(v, (dict, list)):
                    parts.append(f"{self._format_key(k)}: {self._format_complex_value(v)}")
                else:
                    parts.append(f"{self._format_key(k)}: {v}")
        return " | ".join(parts)
    
    def _format_complex_value(self, value: Any) -> str:
        """Format complex values (lists, dicts) as text"""
        if isinstance(value, list):
            if all(isinstance(x, str) for x in value):
                return ", ".join(value)
            return "; ".join(self._format_complex_value(x) if isinstance(x, (dict, list)) else str(x) for x in value)
        elif isinstance(value, dict):
            return ", ".join(f"{k}: {v}" for k, v in value.items() if k not in ['raw_output', 'error'])
        return str(value)
    
    def _add_table(self, doc, headers: List[str], rows: List[List]):
        """Add a formatted table"""
        from docx.shared import Pt, Inches
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for row_data in rows:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row):
                    row[i].text = str(cell_data) if cell_data else ''
        
        doc.add_paragraph()
    
    def _add_section_content(self, doc, section_name: str, content: Any, level: int = 1):
        """Add a section with its content, handling various data types"""
        if content is None:
            return
        
        # Add section heading
        doc.add_heading(self._format_key(section_name), level)
        
        if isinstance(content, str):
            # Plain text content
            self._add_paragraph_text(doc, content)
        
        elif isinstance(content, list):
            if len(content) == 0:
                doc.add_paragraph("No items defined.")
                return
            
            # Check if it's a list of complex objects
            if all(isinstance(item, dict) for item in content):
                self._add_list_of_objects(doc, content, level)
            else:
                self._add_bullet_list(doc, content)
        
        elif isinstance(content, dict):
            self._add_dict_content(doc, content, level)
        
        else:
            doc.add_paragraph(str(content))
    
    def _add_list_of_objects(self, doc, items: List[Dict], parent_level: int):
        """Add a list of objects, each as a subsection"""
        for idx, item in enumerate(items, 1):
            # Try to find a title/name field
            title = item.get('title') or item.get('name') or item.get('id') or item.get('requirement') or f"Item {idx}"
            
            # Add subsection heading
            if parent_level < 3:
                doc.add_heading(f"{idx}. {title}", parent_level + 1)
            else:
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {title}").bold = True
            
            # Add each field
            for key, value in item.items():
                if key in ['title', 'name', 'id', 'raw_output', 'error', 'timestamp']:
                    continue
                
                if value is None or value == '':
                    continue
                
                if isinstance(value, list):
                    p = doc.add_paragraph()
                    p.add_run(f"{self._format_key(key)}:").bold = True
                    self._add_bullet_list(doc, value)
                elif isinstance(value, dict):
                    p = doc.add_paragraph()
                    p.add_run(f"{self._format_key(key)}:").bold = True
                    self._add_dict_content(doc, value, parent_level + 2)
                else:
                    self._add_key_value_paragraph(doc, self._format_key(key), value)
            
            doc.add_paragraph()
    
    def _add_dict_content(self, doc, content: Dict, level: int):
        """Add dictionary content with proper formatting"""
        for key, value in content.items():
            if key in ['raw_output', 'error', 'timestamp']:
                continue
            
            if value is None or value == '':
                continue
            
            if isinstance(value, str):
                self._add_key_value_paragraph(doc, self._format_key(key), value)
            elif isinstance(value, list):
                p = doc.add_paragraph()
                p.add_run(f"{self._format_key(key)}:").bold = True
                if all(isinstance(x, str) for x in value):
                    self._add_bullet_list(doc, value)
                elif all(isinstance(x, dict) for x in value):
                    self._add_list_of_objects(doc, value, level)
                else:
                    self._add_bullet_list(doc, value)
            elif isinstance(value, dict):
                if level < 4:
                    doc.add_heading(self._format_key(key), min(level + 1, 9))
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"{self._format_key(key)}:").bold = True
                self._add_dict_content(doc, value, level + 1)
            else:
                self._add_key_value_paragraph(doc, self._format_key(key), value)
    
    # ==================== BRD CONTENT ====================
    
    def _add_brd_content(self, doc, content: Dict[str, Any]):
        """Add BRD content with proper formatting"""
        sections = [
            ('executive_summary', '1. Executive Summary'),
            ('business_objectives', '2. Business Objectives'),
            ('scope', '3. Scope'),
            ('functional_requirements', '4. Functional Requirements'),
            ('non_functional_requirements', '5. Non-Functional Requirements'),
            ('user_requirements', '6. User Requirements'),
            ('system_requirements', '7. System Requirements'),
            ('data_requirements', '8. Data Requirements'),
            ('constraints_and_assumptions', '9. Constraints and Assumptions'),
            ('appendices', '10. Appendices'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                self._add_brd_section(doc, key, content[key])
                doc.add_page_break()
    
    def _add_brd_section(self, doc, section_key: str, content: Any):
        """Add a specific BRD section"""
        if isinstance(content, str):
            self._add_paragraph_text(doc, content)
            return
        
        if isinstance(content, dict):
            for sub_key, sub_value in content.items():
                if sub_key in ['raw_output', 'error', 'timestamp']:
                    continue
                self._add_section_content(doc, sub_key, sub_value, level=2)
        elif isinstance(content, list):
            self._add_list_of_objects(doc, content, 1)
    
    # ==================== FRD CONTENT ====================
    
    def _add_frd_content(self, doc, content: Dict[str, Any]):
        """Add FRD content with proper formatting"""
        sections = [
            ('introduction', '1. Introduction'),
            ('system_overview', '2. System Overview'),
            ('functional_requirements', '3. Functional Requirements'),
            ('data_requirements', '4. Data Requirements'),
            ('interface_requirements', '5. Interface Requirements'),
            ('security_requirements', '6. Security Requirements'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                self._add_frd_section(doc, key, content[key])
                doc.add_page_break()
    
    def _add_frd_section(self, doc, section_key: str, content: Any):
        """Add a specific FRD section"""
        if isinstance(content, str):
            self._add_paragraph_text(doc, content)
            return
        
        if section_key == 'functional_requirements' and isinstance(content, dict):
            # Handle feature areas structure
            if 'feature_areas' in content:
                for area in content['feature_areas']:
                    doc.add_heading(f"3.x {area.get('area', 'Feature Area')}", 2)
                    if 'description' in area:
                        doc.add_paragraph(area['description'])
                    if 'requirements' in area:
                        self._add_requirements_list(doc, area['requirements'])
            elif 'requirements' in content:
                self._add_requirements_list(doc, content['requirements'])
            else:
                self._add_dict_content(doc, content, 2)
        elif isinstance(content, dict):
            self._add_dict_content(doc, content, 2)
        elif isinstance(content, list):
            self._add_requirements_list(doc, content)
    
    def _add_requirements_list(self, doc, requirements: List):
        """Add a formatted requirements list"""
        for req in requirements:
            if not isinstance(req, dict):
                doc.add_paragraph(str(req), style='List Bullet')
                continue
            
            req_id = req.get('id', 'REQ-XXX')
            title = req.get('title', 'Untitled')
            
            doc.add_heading(f"{req_id}: {title}", 3)
            
            if 'description' in req:
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(req['description'])
            
            if 'rationale' in req:
                p = doc.add_paragraph()
                p.add_run("Rationale: ").bold = True
                p.add_run(req['rationale'])
            
            if 'acceptance_criteria' in req:
                p = doc.add_paragraph()
                p.add_run("Acceptance Criteria:").bold = True
                criteria = req['acceptance_criteria']
                if isinstance(criteria, list):
                    for ac in criteria:
                        doc.add_paragraph(str(ac), style='List Bullet')
                else:
                    doc.add_paragraph(str(criteria))
            
            if 'priority' in req:
                p = doc.add_paragraph()
                p.add_run("Priority: ").bold = True
                p.add_run(req['priority'])
            
            if 'source' in req:
                p = doc.add_paragraph()
                p.add_run("Source: ").bold = True
                p.add_run(req['source'])
            
            if 'dependencies' in req:
                p = doc.add_paragraph()
                p.add_run("Dependencies: ").bold = True
                deps = req['dependencies']
                p.add_run(", ".join(deps) if isinstance(deps, list) else str(deps))
            
            doc.add_paragraph()
    
    # ==================== USER STORIES CONTENT ====================
    
    def _add_user_stories_content(self, doc, content: Any):
        """Add User Stories content with proper formatting"""
        doc.add_heading('1. User Stories Overview', 1)
        
        # Handle epics if present
        if isinstance(content, dict) and 'epics_overview' in content:
            doc.add_heading('1.1 Epics', 2)
            epics = content['epics_overview']
            if isinstance(epics, dict) and 'epics' in epics:
                epics = epics['epics']
            if isinstance(epics, list):
                for epic in epics:
                    if isinstance(epic, dict):
                        doc.add_heading(f"{epic.get('id', 'E-XXX')}: {epic.get('title', 'Epic')}", 3)
                        if 'description' in epic:
                            doc.add_paragraph(epic['description'])
                        if 'business_value' in epic:
                            p = doc.add_paragraph()
                            p.add_run("Business Value: ").bold = True
                            p.add_run(epic['business_value'])
        
        doc.add_page_break()
        
        # Collect all stories
        all_stories = []
        
        if isinstance(content, list):
            all_stories = content
        elif isinstance(content, dict):
            if 'stories' in content:
                all_stories = content['stories']
            else:
                # Check for story sets
                for key in ['user_stories_set1', 'user_stories_set2']:
                    if key in content:
                        stories = content[key]
                        if isinstance(stories, list):
                            all_stories.extend(stories)
                        elif isinstance(stories, dict) and 'stories' in stories:
                            all_stories.extend(stories['stories'])
        
        if all_stories:
            doc.add_heading('2. User Stories', 1)
            self._add_user_stories_list(doc, all_stories)
        
        # Story mapping if present
        if isinstance(content, dict) and 'story_mapping' in content:
            doc.add_page_break()
            doc.add_heading('3. Story Mapping', 1)
            self._add_section_content(doc, 'story_mapping', content['story_mapping'], 2)
    
    def _add_user_stories_list(self, doc, stories: List):
        """Add formatted user stories"""
        for idx, story in enumerate(stories, 1):
            if not isinstance(story, dict):
                continue
            
            story_id = story.get('id', f'US-{idx:03d}')
            title = story.get('title', 'Untitled Story')
            
            doc.add_heading(f"{story_id}: {title}", 2)
            
            if 'epic' in story or 'epic_reference' in story:
                p = doc.add_paragraph()
                p.add_run("Epic: ").bold = True
                p.add_run(story.get('epic') or story.get('epic_reference', 'N/A'))
            
            if 'story' in story:
                p = doc.add_paragraph()
                p.add_run("User Story: ").bold = True
                p.add_run(story['story'])
            
            if 'description' in story:
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(story['description'])
            
            if 'acceptance_criteria' in story:
                p = doc.add_paragraph()
                p.add_run("Acceptance Criteria:").bold = True
                criteria = story['acceptance_criteria']
                if isinstance(criteria, list):
                    for ac in criteria:
                        doc.add_paragraph(str(ac), style='List Bullet')
                else:
                    doc.add_paragraph(str(criteria))
            
            # Metadata row
            meta_parts = []
            if 'priority' in story:
                meta_parts.append(f"Priority: {story['priority']}")
            if 'story_points' in story:
                meta_parts.append(f"Story Points: {story['story_points']}")
            if 'dependencies' in story:
                deps = story['dependencies']
                if isinstance(deps, list):
                    meta_parts.append(f"Dependencies: {', '.join(deps)}")
            
            if meta_parts:
                p = doc.add_paragraph()
                p.add_run(" | ".join(meta_parts)).italic = True
            
            doc.add_paragraph()
    
    # ==================== TEST CASES CONTENT ====================
    
    def _add_test_cases_content(self, doc, content: Any):
        """Add Test Cases content with proper formatting"""
        # Test Strategy section
        if isinstance(content, dict) and 'test_strategy' in content:
            doc.add_heading('1. Test Strategy', 1)
            self._add_section_content(doc, 'test_strategy', content['test_strategy'], 2)
            doc.add_page_break()
        
        # Collect all test cases
        all_tests = []
        section_num = 2
        
        if isinstance(content, list):
            all_tests = [('Test Cases', content)]
        elif isinstance(content, dict):
            if 'test_cases' in content:
                all_tests = [('Test Cases', content['test_cases'])]
            else:
                test_sections = [
                    ('functional_test_cases', 'Functional Test Cases'),
                    ('negative_test_cases', 'Negative Test Cases'),
                    ('integration_test_cases', 'Integration Test Cases'),
                    ('performance_test_cases', 'Performance Test Cases'),
                ]
                for key, title in test_sections:
                    if key in content:
                        tests = content[key]
                        if isinstance(tests, dict) and 'test_cases' in tests:
                            tests = tests['test_cases']
                        if isinstance(tests, list):
                            all_tests.append((title, tests))
        
        for title, tests in all_tests:
            doc.add_heading(f'{section_num}. {title}', 1)
            self._add_test_cases_list(doc, tests)
            doc.add_page_break()
            section_num += 1
    
    def _add_test_cases_list(self, doc, test_cases: List):
        """Add formatted test cases"""
        for idx, tc in enumerate(test_cases, 1):
            if not isinstance(tc, dict):
                continue
            
            tc_id = tc.get('id', tc.get('test_id', f'TC-{idx:03d}'))
            name = tc.get('name', tc.get('title', 'Untitled Test'))
            
            doc.add_heading(f"{tc_id}: {name}", 2)
            
            if 'description' in tc:
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(tc['description'])
            
            if 'preconditions' in tc:
                p = doc.add_paragraph()
                p.add_run("Preconditions: ").bold = True
                precond = tc['preconditions']
                if isinstance(precond, list):
                    p.add_run("; ".join(precond))
                else:
                    p.add_run(str(precond))
            
            if 'test_data' in tc:
                p = doc.add_paragraph()
                p.add_run("Test Data: ").bold = True
                p.add_run(str(tc['test_data']))
            
            # Test Steps
            steps = tc.get('steps') or tc.get('test_steps', [])
            if steps:
                p = doc.add_paragraph()
                p.add_run("Test Steps:").bold = True
                if isinstance(steps, list):
                    for i, step in enumerate(steps, 1):
                        doc.add_paragraph(f"{i}. {step}", style='List Number')
                else:
                    doc.add_paragraph(str(steps))
            
            # Expected Results
            expected = tc.get('expected_results') or tc.get('expected', '')
            if expected:
                p = doc.add_paragraph()
                p.add_run("Expected Results: ").bold = True
                if isinstance(expected, list):
                    p.add_run("; ".join(expected))
                else:
                    p.add_run(str(expected))
            
            # Metadata
            meta_parts = []
            if 'priority' in tc:
                meta_parts.append(f"Priority: {tc['priority']}")
            if 'test_type' in tc:
                meta_parts.append(f"Type: {tc['test_type']}")
            if 'category' in tc:
                meta_parts.append(f"Category: {tc['category']}")
            
            if meta_parts:
                p = doc.add_paragraph()
                p.add_run(" | ".join(meta_parts)).italic = True
            
            doc.add_paragraph()
    
    # ==================== MIGRATION PLAN CONTENT ====================
    
    def _add_migration_plan_content(self, doc, content: Dict[str, Any]):
        """Add Migration Plan content"""
        sections = [
            ('executive_summary', '1. Executive Summary'),
            ('current_state_analysis', '2. Current State Analysis'),
            ('target_state_design', '3. Target State Design'),
            ('migration_phases', '4. Migration Phases'),
            ('risk_and_rollback', '5. Risk Management and Rollback'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                section_content = content[key]
                
                if key == 'migration_phases' and isinstance(section_content, list):
                    self._add_migration_phases(doc, section_content)
                else:
                    self._add_section_content(doc, key, section_content, 2)
                
                doc.add_page_break()
    
    def _add_migration_phases(self, doc, phases: List):
        """Add migration phases with detailed formatting"""
        for idx, phase in enumerate(phases, 1):
            if not isinstance(phase, dict):
                continue
            
            phase_name = phase.get('name', phase.get('phase', f'Phase {idx}'))
            doc.add_heading(f"4.{idx} {phase_name}", 2)
            
            if 'objectives' in phase:
                p = doc.add_paragraph()
                p.add_run("Objectives: ").bold = True
                p.add_run(str(phase['objectives']))
            
            if 'duration' in phase:
                p = doc.add_paragraph()
                p.add_run("Duration: ").bold = True
                p.add_run(str(phase['duration']))
            
            if 'activities' in phase:
                p = doc.add_paragraph()
                p.add_run("Activities:").bold = True
                activities = phase['activities']
                if isinstance(activities, list):
                    for act in activities:
                        doc.add_paragraph(str(act), style='List Bullet')
            
            if 'deliverables' in phase:
                p = doc.add_paragraph()
                p.add_run("Deliverables:").bold = True
                deliverables = phase['deliverables']
                if isinstance(deliverables, list):
                    for d in deliverables:
                        doc.add_paragraph(str(d), style='List Bullet')
            
            if 'resources' in phase:
                p = doc.add_paragraph()
                p.add_run("Resources Required: ").bold = True
                p.add_run(str(phase['resources']))
            
            if 'risks' in phase:
                p = doc.add_paragraph()
                p.add_run("Risks:").bold = True
                risks = phase['risks']
                if isinstance(risks, list):
                    for r in risks:
                        doc.add_paragraph(str(r), style='List Bullet')
            
            if 'success_criteria' in phase:
                p = doc.add_paragraph()
                p.add_run("Success Criteria: ").bold = True
                criteria = phase['success_criteria']
                if isinstance(criteria, list):
                    p.add_run("; ".join(criteria))
                else:
                    p.add_run(str(criteria))
            
            doc.add_paragraph()
    
    # ==================== REVERSE ENGINEERING CONTENT ====================
    
    def _add_reverse_eng_content(self, doc, content: Dict[str, Any]):
        """Add Reverse Engineering document content"""
        sections = [
            ('system_overview', '1. System Overview'),
            ('architecture_analysis', '2. Architecture Analysis'),
            ('technology_stack', '3. Technology Stack'),
            ('data_model', '4. Data Model'),
            ('technical_debt', '5. Technical Debt Assessment'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                self._add_section_content(doc, key, content[key], 2)
                doc.add_page_break()
    
    # ==================== TDD CONTENT ====================
    
    def _add_tdd_content(self, doc, content: Dict[str, Any]):
        """Add Technical Design Document content"""
        sections = [
            ('technical_overview', '1. Technical Overview'),
            ('architecture_design', '2. Architecture Design'),
            ('data_design', '3. Data Design'),
            ('api_design', '4. API Design'),
            ('security_design', '5. Security Design'),
            ('deployment_design', '6. Deployment Design'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                self._add_section_content(doc, key, content[key], 2)
                doc.add_page_break()
    
    # ==================== DB ANALYSIS CONTENT ====================
    
    def _add_db_analysis_content(self, doc, content: Dict[str, Any]):
        """Add Database Analysis content"""
        sections = [
            ('database_overview', '1. Database Overview'),
            ('schema_analysis', '2. Schema Analysis'),
            ('index_analysis', '3. Index Analysis'),
            ('optimization_recommendations', '4. Optimization Recommendations'),
        ]
        
        for key, title in sections:
            if key in content:
                doc.add_heading(title, 1)
                self._add_section_content(doc, key, content[key], 2)
                doc.add_page_break()
    
    # ==================== GENERIC CONTENT ====================
    
    def _add_generic_content(self, doc, content: Any):
        """Add generic content for unknown types"""
        if isinstance(content, dict):
            for key, value in content.items():
                if key not in ['raw_output', 'error', 'timestamp']:
                    self._add_section_content(doc, key, value, 1)
        elif isinstance(content, list):
            doc.add_heading('Content', 1)
            self._add_bullet_list(doc, content)
        else:
            doc.add_paragraph(str(content))
    
    # ==================== MARKDOWN EXPORT ====================
    
    def to_markdown(self, content: Any, analysis_type: str) -> bytes:
        """Export document to Markdown format"""
        try:
            md = []
            md.append(f"# {self._get_document_title(analysis_type)}\n")
            md.append(f"**Document Type:** {analysis_type.upper()}\n")
            md.append("---\n")
            
            content_dict = self._parse_content(content)
            md.append(self._content_to_markdown(content_dict, 2))
            
            return '\n'.join(md).encode('utf-8')
        except Exception as e:
            logger.error(f"Error creating Markdown: {str(e)}")
            return self._export_as_text(content, analysis_type).encode('utf-8')
    
    def _content_to_markdown(self, content: Any, level: int = 2) -> str:
        """Convert content to Markdown"""
        md = []
        
        if isinstance(content, dict):
            for key, value in content.items():
                if key in ['raw_output', 'error', 'timestamp']:
                    continue
                
                title = self._format_key(key)
                md.append(f"\n{'#' * min(level, 6)} {title}\n")
                md.append(self._content_to_markdown(value, level + 1))
        
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    # Get item title
                    title = item.get('title') or item.get('name') or item.get('id') or 'Item'
                    md.append(f"\n### {title}\n")
                    for k, v in item.items():
                        if k not in ['title', 'name', 'id', 'raw_output']:
                            if isinstance(v, list):
                                md.append(f"**{self._format_key(k)}:**")
                                for vi in v:
                                    md.append(f"- {vi}")
                            else:
                                md.append(f"**{self._format_key(k)}:** {v}")
                    md.append("")
                else:
                    md.append(f"- {item}")
        
        else:
            md.append(str(content))
        
        return '\n'.join(md)
    
    # ==================== HTML EXPORT ====================
    
    def to_html(self, content: Any, analysis_type: str, project_id: str) -> bytes:
        """Export document to HTML format"""
        try:
            content_dict = self._parse_content(content)
            
            html = [
                '<!DOCTYPE html>',
                '<html>',
                '<head>',
                '<meta charset="UTF-8">',
                f'<title>{self._get_document_title(analysis_type)}</title>',
                '<style>',
                self._get_html_styles(),
                '</style>',
                '</head>',
                '<body>',
                '<div class="container">',
                f'<h1>{self._get_document_title(analysis_type)}</h1>',
                f'<p class="meta">Project ID: {project_id}</p>',
                f'<p class="meta">Document Type: {analysis_type.upper()}</p>',
                '<hr>',
                self._content_to_html(content_dict, 2),
                '</div>',
                '</body>',
                '</html>'
            ]
            
            return '\n'.join(html).encode('utf-8')
        except Exception as e:
            logger.error(f"Error creating HTML: {str(e)}")
            return f"<html><body><pre>{self._export_as_text(content, analysis_type)}</pre></body></html>".encode('utf-8')
    
    def _content_to_html(self, content: Any, level: int = 2) -> str:
        """Convert content to HTML"""
        html = []
        
        if isinstance(content, dict):
            for key, value in content.items():
                if key in ['raw_output', 'error', 'timestamp']:
                    continue
                
                title = self._format_key(key)
                html.append(f'<h{min(level, 6)}>{title}</h{min(level, 6)}>')
                html.append(self._content_to_html(value, level + 1))
        
        elif isinstance(content, list):
            html.append('<div class="list-section">')
            for item in content:
                if isinstance(item, dict):
                    title = item.get('title') or item.get('name') or item.get('id') or 'Item'
                    html.append(f'<div class="item"><h{min(level+1, 6)}>{title}</h{min(level+1, 6)}>')
                    for k, v in item.items():
                        if k not in ['title', 'name', 'id', 'raw_output']:
                            if isinstance(v, list):
                                html.append(f'<p><strong>{self._format_key(k)}:</strong></p>')
                                html.append('<ul>')
                                for vi in v:
                                    html.append(f'<li>{vi}</li>')
                                html.append('</ul>')
                            else:
                                html.append(f'<p><strong>{self._format_key(k)}:</strong> {v}</p>')
                    html.append('</div>')
                else:
                    html.append(f'<li>{item}</li>')
            html.append('</div>')
        
        else:
            html.append(f'<p>{content}</p>')
        
        return '\n'.join(html)
    
    def _get_html_styles(self) -> str:
        """Get CSS styles for HTML export"""
        return """
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }
        .container {
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
        h2 { color: #34495e; margin-top: 30px; border-bottom: 1px solid #eee; padding-bottom: 5px; }
        h3 { color: #7f8c8d; }
        .meta { color: #7f8c8d; font-size: 0.9em; }
        .item { border-left: 4px solid #3498db; padding-left: 20px; margin: 20px 0; }
        ul, ol { margin: 10px 0; }
        hr { border: none; border-top: 1px solid #ecf0f1; margin: 20px 0; }
        .list-section { margin: 15px 0; }
        """
    
    # ==================== HELPER METHODS ====================
    
    def _get_document_title(self, analysis_type: str) -> str:
        """Get formatted document title"""
        titles = {
            'brd': 'Business Requirements Document',
            'frd': 'Functional Requirements Document',
            'user_stories': 'User Stories Document',
            'test_cases': 'Test Cases Document',
            'migration_plan': 'Migration Plan',
            'reverse_eng': 'Reverse Engineering Document',
            'tdd': 'Technical Design Document',
            'db_analysis': 'Database Analysis Document'
        }
        return titles.get(analysis_type, analysis_type.upper())
    
    def _export_as_text(self, content: Any, analysis_type: str) -> str:
        """Export as plain text (fallback)"""
        text = [
            self._get_document_title(analysis_type),
            '=' * 50,
            '',
        ]
        
        if isinstance(content, dict):
            for key, value in content.items():
                text.append(f"\n## {self._format_key(key)}\n")
                text.append(self._format_complex_value(value))
        else:
            text.append(json.dumps(content, indent=2))
        
        return '\n'.join(text)
